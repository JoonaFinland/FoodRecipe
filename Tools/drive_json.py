import os
import json
import docx
import string

"""
def parseJSON(file):
    file_path = os.path.join(BASE, JSON, file)
    with open(file_path, 'r', encoding="utf-8") as f:
        data = json.load(f)
        
    for food in data['foods']:
        print('Name: '+food['name'])
        for ing in food['ings']:
            print('Sub ingredient: '+ing['name'])
            print(ing['ings'].split('|'))
            
        print('Steps: ')
        print(food['steps'].split('|'))
        print('Hints: ')
        print(food['hints'].split('|'))
'''
text_entries = os.listdir(os.path.join(BASE, JSON))

for text in text_entries:
    parseJSON(text)'''
    
    
with open(os.path.join(BASE, JSON, "BBQ-kanapizza.json"), 'r', encoding='utf-8') as f:
    data = json.load(f)
    print(json.dumps(data['foods'][0]['ings']))"""
    
    
def new_obj(curr, prev):
    return prev in ['foods','hints'] and curr == 'ings'
    
    
def appendDict(dict,key,val):
    if key in dict:
        dict[key].append(val)
    else:
        dict[key] = [val]
    return dict

BASE = "KEITTOKIRJA"
WORD = "Reseptit"
JSON = "Json"
TMP = ''
IMGS = "Reseptien kuvat"
# Chunk types
STEP = "OHJE:"
HINT = "VINKKI:"
    
def doc_to_json(file):
    doc = docx.Document(os.path.join(BASE,WORD,file))

    title_raw = doc.paragraphs[0].text
    name = title_raw.split('(')[0].rstrip().lower().capitalize()
    portions = title_raw.split('(')[-1].split(' ')[0]
    check_chunk = True
    prev_chunk = ''

    foods = []

    tmp_obj = {
        "name": name,
        "portions": portions,
        "ings": {},
        "steps": [],
        "hints": [],
    }
    for para in doc.paragraphs[2:]:
        if para.text == '':
            check_chunk = True
            continue
        if check_chunk:
            check_chunk = False
            if STEP in para.text:
                cur_chunk = "step"
            elif HINT in para.text:
                cur_chunk = "hint"
            else:
                cur_chunk = "ings"
                ings_curr = para.text.replace(':','').lower().capitalize().rstrip()
            if new_obj(cur_chunk,prev_chunk):
                # new obj
                foods.append(tmp_obj)
                tmp_obj = {
                    "name": '',
                    "portions": portions,
                    "ings": [],
                    "steps": '',
                    "hints": '',
                }
            continue
        if cur_chunk == "ings":
            ing_row = para.text
            if ing_row.startswith('n.'):
                ing_out = ' '.join(ing_row.split(' ')[:3])+'+'+' '.join(ing_row.split(' ')[3:])
            elif ing_row[0] not in string.ascii_lowercase:
                ing_out = ' '.join(ing_row.split(' ')[:2])+'+'+' '.join(ing_row.split(' ')[2:])
            else:
                ing_out = '+'+ing_row
            tmp_obj['ings'] = appendDict(tmp_obj['ings'],ings_curr,ing_out.rstrip())
        elif cur_chunk == "step":
            tmp_obj['steps'].append(para.text.rstrip())
        elif cur_chunk == "hint":
            tmp_obj['hints'].append(para.text.rstrip())
        else:
            print('ERROR: BAD cur_chunk: {}'.format(cur_chunk))
            
        
    foods.append(tmp_obj)

    main = {'foods':foods}


    #print(json.dumps(main, indent=4))

    with open(os.path.join(BASE,TMP,file.split('.')[0]+'.json'),'w',encoding='utf-8') as f:
        json.dump(main,f, indent=4, ensure_ascii=False)
    
    return '{} saved successfully'.format(file.split('.')[0]+'.json')


text_entries = os.listdir(os.path.join(BASE, WORD))

for text in text_entries:
    print(doc_to_json(text))